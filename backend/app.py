# from flask import Flask, request, jsonify
# from flask_cors import CORS
# import os
# from werkzeug.utils import secure_filename
# import logging
# from datetime import datetime
# import uuid
# import json
# import csv


# # Configure logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# app = Flask(__name__)
# CORS(app)  # Enable CORS for all routes

# # Configuration
# UPLOAD_FOLDER = 'uploads'
# MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max upload size
# ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'csv', 'xls', 'xlsx'}

# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# # Create upload directory if it doesn't exist
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# # In-memory storage for chat history and documents
# chat_history = {}
# uploaded_documents = {}

# # Helper functions
# def allowed_file(filename):
#     """Check if the file is allowed based on its extension"""
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_file(file_path):
#     """Extract text content from different file types"""
#     file_extension = file_path.rsplit('.', 1)[1].lower()

#     if file_extension == 'txt':
#         with open(file_path, 'r', encoding='utf-8') as file:
#             return file.read()
    
#     elif file_extension == 'pdf':
#         with open(file_path, 'rb') as file:
#             reader = PyPDF2.PdfReader(file)
#             text = ""
#             for page in reader.pages:
#                 text += page.extract_text()
#             return text
    
#     elif file_extension == 'docx':
#         doc = docx.Document(file_path)
#         text = ""
#         for para in doc.paragraphs:
#             text += para.text + '\n'
#         return text
    
#     elif file_extension in ['csv', 'xls', 'xlsx']:
#         text = ""
#         with open(file_path, 'r') as file:
#             if file_extension == 'csv':
#                 reader = csv.reader(file)
#                 for row in reader:
#                     text += ', '.join(row) + '\n'
#             elif file_extension in ['xls', 'xlsx']:
#                 wb = xlrd.open_workbook(file_path)
#                 sheet = wb.sheet_by_index(0)
#                 for row in range(sheet.nrows):
#                     text += ' | '.join(str(sheet.cell_value(row, col)) for col in range(sheet.ncols)) + '\n'
#         return text

# @app.route('/process_query', methods=['POST'])
# def process_query():
#     try:
#         # Parse the incoming JSON data
#         data = request.get_json()
#         query = data.get('query', '').lower()  # Normalize the query to lowercase

#         # Check if the query type is related to a company
#         if "company" in query:
#             company_name = query.split("company")[-1].strip()  # Extract company name from the query
#             response_data = handle_company_query(company_name)
        
#         # Check if the query is related to a sector
#         elif "sector" in query:
#             sector_name = query.split("sector")[-1].strip()  # Extract sector name from the query
#             response_data = handle_sector_query(sector_name)

#         # Default response for unrecognized queries
#         else:
#             response_data = {
#                 "type": "unknown",
#                 "message": "Could not identify the query type."
#             }

#         return jsonify({"response": json.dumps(response_data)})

#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# def process_document(text_content):
#     """Process document content to generate a summary"""
#     # Placeholder: You can add text summarization or any LLM processing here
#     return f"Summary: {text_content[:100]}..."  # Just a placeholder summary

# @app.route('/api/chat', methods=['POST'])
# def chat():
#     """Process chat messages and return AI responses"""
#     try:
#         data = request.json
#         if not data or 'message' not in data:
#             return jsonify({'error': 'No message provided'}), 400
        
#         message = data['message']
#         user_id = data.get('user_id', 'default_user')
        
#         if user_id not in chat_history:
#             chat_history[user_id] = []
        
#         # Add user message to history
#         chat_history[user_id].append({
#             'role': 'user',
#             'content': message,
#             'timestamp': datetime.now().isoformat()
#         })
        
#         # Process with LLM (or placeholder)
#         context = []
#         if user_id in uploaded_documents:
#             for doc_id, doc_info in uploaded_documents[user_id].items():
#                 context.append(doc_info['content'])
        
#         response = process_query(message, context, chat_history[user_id])
        
#         # Add assistant response to history
#         chat_history[user_id].append({
#             'role': 'assistant',
#             'content': response,
#             'timestamp': datetime.now().isoformat()
#         })
        
#         return jsonify({
#             'response': response,
#             'timestamp': datetime.now().isoformat()
#         })
    
#     except Exception as e:
#         logger.error(f"Error in chat endpoint: {str(e)}")
#         return jsonify({'error': str(e)}), 500

# @app.route('/api/upload', methods=['POST'])
# def upload_file():
#     """Handle file uploads and process documents"""
#     try:
#         if 'files' not in request.files:
#             return jsonify({'error': 'No file part'}), 400
        
#         files = request.files.getlist('files')
#         user_id = request.form.get('user_id', 'default_user')
        
#         if user_id not in uploaded_documents:
#             uploaded_documents[user_id] = {}
        
#         results = []
        
#         for file in files:
#             if file.filename == '':
#                 continue
                
#             if file and allowed_file(file.filename):
#                 filename = secure_filename(file.filename)
#                 file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#                 file.save(file_path)
                
#                 # Extract text from file
#                 text_content = extract_text_from_file(file_path)
                
#                 # Process document with LLM (or placeholder)
#                 summary = process_document(text_content)
                
#                 # Generate a unique ID for this document
#                 doc_id = str(uuid.uuid4())
                
#                 # Store document info
#                 uploaded_documents[user_id][doc_id] = {
#                     'filename': filename,
#                     'path': file_path,
#                     'content': text_content,
#                     'summary': summary,
#                     'upload_time': datetime.now().isoformat()
#                 }
                
#                 results.append({
#                     'doc_id': doc_id,
#                     'filename': filename,
#                     'summary': summary
#                 })
        
#         return jsonify({
#             'message': f'Successfully processed {len(results)} file(s)',
#             'results': results
#         })
    
#     except Exception as e:
#         logger.error(f"Error in upload endpoint: {str(e)}")
#         return jsonify({'error': str(e)}), 500

# @app.route('/api/documents', methods=['GET'])
# def get_documents():
#     """Get list of uploaded documents for a user"""
#     user_id = request.args.get('user_id', 'default_user')
    
#     if user_id not in uploaded_documents:
#         return jsonify({'documents': []})
    
#     docs = []
#     for doc_id, doc_info in uploaded_documents[user_id].items():
#         docs.append({
#             'doc_id': doc_id,
#             'filename': doc_info['filename'],
#             'summary': doc_info['summary'],
#             'upload_time': doc_info['upload_time']
#         })
    
#     return jsonify({'documents': docs})

# @app.route('/api/reports/<report_id>', methods=['GET'])
# def get_report(report_id):
#     """Get a specific report"""
#     # Placeholder for fetching report data from database
#     report = {
#         'id': report_id,
#         'title': 'Sample Report',
#         'sections': [
#             {
#                 'title': 'Overview',
#                 'content': 'This is an overview section of the report.'
#             },
#             {
#                 'title': 'Key Metrics',
#                 'content': 'This section contains key metrics and analysis.'
#             },
#             {
#                 'title': 'Recommendations',
#                 'content': 'Based on the analysis, here are the recommendations.'
#             }
#         ],
#         'created_at': datetime.now().isoformat()
#     }
    
#     return jsonify(report)

# @app.route('/api/health', methods=['GET'])
# def health_check():
#     """Health check endpoint"""
#     return jsonify({'status': 'ok', 'version': '1.0.0'})

# if __name__ == '__main__':
#     app.run(debug=True, host='0.0.0.0', port=5000)

from flask import Flask, request, jsonify, render_template, send_from_directory, send_file
from flask_cors import CORS
import os
from werkzeug.utils import secure_filename
import logging
from datetime import datetime
import uuid
import json
import csv
import PyPDF2
import docx
import xlrd

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__, static_folder='../public', static_url_path='/')
CORS(app)  # Enable CORS for all routes - essential for Next.js frontend

# Configuration
UPLOAD_FOLDER = 'uploads'
TEMPLATE_FOLDER = 'templates'
STATIC_FOLDER = 'static'
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max upload size
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'csv', 'xls', 'xlsx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Create necessary directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEMPLATE_FOLDER, exist_ok=True)
os.makedirs(STATIC_FOLDER, exist_ok=True)

# In-memory storage for chat history and documents
chat_history = {}
uploaded_documents = {}

# Helper functions
def allowed_file(filename):
    """Check if the file is allowed based on its extension"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path):
    """Extract text content from different file types"""
    try:
        logger.info(f"Extracting text from: {file_path}")
        file_extension = file_path.rsplit('.', 1)[1].lower()

        if file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        
        elif file_extension == 'pdf':
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() or ""  # Handle None returns
                    return text
            except Exception as e:
                logger.error(f"Error extracting PDF text: {str(e)}")
                return f"Error extracting PDF: {str(e)}"
        
        elif file_extension == 'docx':
            try:
                doc = docx.Document(file_path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + '\n'
                return text
            except Exception as e:
                logger.error(f"Error extracting DOCX text: {str(e)}")
                return f"Error extracting DOCX: {str(e)}"
        
        elif file_extension in ['csv', 'xls', 'xlsx']:
            text = ""
            try:
                if file_extension == 'csv':
                    with open(file_path, 'r', encoding='utf-8') as file:
                        reader = csv.reader(file)
                        for row in reader:
                            text += ', '.join(row) + '\n'
                elif file_extension in ['xls', 'xlsx']:
                    wb = xlrd.open_workbook(file_path)
                    sheet = wb.sheet_by_index(0)
                    for row in range(sheet.nrows):
                        text += ' | '.join(str(sheet.cell_value(row, col)) for col in range(sheet.ncols)) + '\n'
                return text
            except Exception as e:
                logger.error(f"Error extracting spreadsheet text: {str(e)}")
                return f"Error extracting spreadsheet: {str(e)}"
        
        logger.info(f"Successfully extracted text from {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from file {file_path}: {str(e)}")
        return f"Error extracting text: {str(e)}"

def handle_company_query(company_name):
    """Handle company-related queries"""
    logger.info(f"Processing company query for: {company_name}")
    # Placeholder implementation
    return {
        "type": "company",
        "company_name": company_name,
        "message": f"Information about {company_name}"
    }

def handle_sector_query(sector_name):
    """Handle sector-related queries"""
    logger.info(f"Processing sector query for: {sector_name}")
    # Placeholder implementation
    return {
        "type": "sector",
        "sector_name": sector_name,
        "message": f"Information about {sector_name} sector"
    }

def process_query(message, context, history):
    """Process a query message and generate a response"""
    logger.info(f"Processing query: {message}")
    # Placeholder for actual LLM processing
    response = f"Received: {message}"
    
    if context:
        response += f"\nWith {len(context)} document(s) as context"
    
    if history and len(history) > 1:
        response += f"\nBased on conversation with {len(history)//2} exchanges"
    
    logger.info(f"Generated response for query")
    return response

def process_document(text_content):
    """Process document content to generate a summary"""
    logger.info("Processing document content for summary")
    # Placeholder: You can add text summarization or any LLM processing here
    summary = text_content[:100] + "..." if len(text_content) > 100 else text_content
    return f"Summary: {summary}"  # Just a placeholder summary

# API Routes for Next.js frontend
@app.route('/api/chat', methods=['POST'])
def chat():
    """Process chat messages and return AI responses"""
    try:
        logger.info("Received request to /api/chat endpoint")
        data = request.json
        if not data or 'message' not in data:
            logger.warning("No message provided in chat request")
            return jsonify({'error': 'No message provided'}), 400
        
        message = data['message']
        user_id = data.get('user_id', 'default_user')
        logger.info(f"Chat request from user {user_id}: {message}")
        
        # Handle user type for different processing levels
        user_type = data.get('user_type', 'normal-user')
        logger.info(f"User type: {user_type}")
        
        if user_id not in chat_history:
            logger.info(f"Creating new chat history for user {user_id}")
            chat_history[user_id] = []
        
        # Add user message to history
        chat_history[user_id].append({
            'role': 'user',
            'content': message,
            'timestamp': datetime.now().isoformat()
        })
        
        # Process with LLM (or placeholder)
        context = []
        if user_id in uploaded_documents:
            for doc_id, doc_info in uploaded_documents[user_id].items():
                context.append(doc_info['content'])
            logger.info(f"Using {len(context)} documents as context for user {user_id}")
        
        response = process_query(message, context, chat_history[user_id])
        
        # If professional user, add more detailed analysis
        if user_type == 'professional-user':
            response += "\n\nAdditional professional insights: This analysis includes deeper financial metrics and forecasting data that would be valuable for professional users."
        
        # Add assistant response to history
        chat_history[user_id].append({
            'role': 'assistant',
            'content': response,
            'timestamp': datetime.now().isoformat()
        })
        
        logger.info(f"Generated response for user {user_id}")
        return jsonify({
            'response': response,
            'timestamp': datetime.now().isoformat()
        })
    
    except Exception as e:
        logger.error(f"Error in chat endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file uploads and process documents"""
    try:
        logger.info("Received request to /api/upload endpoint")
        if 'files' not in request.files:
            logger.warning("No files part in the request")
            return jsonify({'error': 'No file part'}), 400
        
        files = request.files.getlist('files')
        if not files or files[0].filename == '':
            logger.warning("No files selected")
            return jsonify({'error': 'No files selected'}), 400
            
        user_id = request.form.get('user_id', 'default_user')
        user_type = request.form.get('user_type', 'normal-user')
        logger.info(f"Processing {len(files)} file(s) for {user_type} - user {user_id}")
        
        if user_id not in uploaded_documents:
            uploaded_documents[user_id] = {}
        
        results = []
        
        for file in files:
            if file.filename == '':
                logger.warning("Empty filename submitted")
                continue
                
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                unique_filename = f"{uuid.uuid4()}_{filename}"  # Ensure unique filenames
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                logger.info(f"Saving file {filename} to {file_path}")
                file.save(file_path)
                
                # Extract text from file
                text_content = extract_text_from_file(file_path)
                
                # Process document with LLM (or placeholder)
                summary = process_document(text_content)
                
                # For professional users, add more detailed analysis
                if user_type == 'professional-user':
                    summary += " (Professional analysis: Additional financial metrics and insights provided)"
                
                # Generate a unique ID for this document
                doc_id = str(uuid.uuid4())
                logger.info(f"Assigned document ID {doc_id} to {filename}")
                
                # Store document info
                uploaded_documents[user_id][doc_id] = {
                    'filename': filename,
                    'path': file_path,
                    'content': text_content,
                    'summary': summary,
                    'user_type': user_type,
                    'upload_time': datetime.now().isoformat()
                }
                
                results.append({
                    'doc_id': doc_id,
                    'filename': filename,
                    'summary': summary
                })
            else:
                logger.warning(f"File {file.filename} has a disallowed extension")
                results.append({
                    'filename': file.filename,
                    'error': 'File type not allowed'
                })
        
        if not results:
            return jsonify({
                'message': 'No valid files were processed',
                'results': []
            }), 400
            
        logger.info(f"Successfully processed {len(results)} file(s)")
        return jsonify({
            'message': f'Successfully processed {len(results)} file(s)',
            'results': results
        })
    
    except Exception as e:
        logger.error(f"Error in upload endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get list of uploaded documents for a user"""
    try:
        user_id = request.args.get('user_id', 'default_user')
        user_type = request.args.get('user_type', None)  # Filter by user type if specified
        logger.info(f"Getting documents for user {user_id}, type filter: {user_type}")
        
        if user_id not in uploaded_documents:
            logger.info(f"No documents found for user {user_id}")
            return jsonify({'documents': []})
        
        docs = []
        for doc_id, doc_info in uploaded_documents[user_id].items():
            # Filter by user type if specified
            if user_type and doc_info.get('user_type') != user_type:
                continue
                
            docs.append({
                'doc_id': doc_id,
                'filename': doc_info['filename'],
                'summary': doc_info['summary'],
                'user_type': doc_info.get('user_type', 'normal-user'),
                'upload_time': doc_info['upload_time']
            })
        
        logger.info(f"Returning {len(docs)} documents for user {user_id}")
        return jsonify({'documents': docs})
    except Exception as e:
        logger.error(f"Error in get_documents endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents/<doc_id>', methods=['GET'])
def get_document(doc_id):
    """Get a specific document by ID"""
    try:
        user_id = request.args.get('user_id', 'default_user')
        logger.info(f"Retrieving document {doc_id} for user {user_id}")
        
        if user_id not in uploaded_documents or doc_id not in uploaded_documents[user_id]:
            logger.warning(f"Document {doc_id} not found for user {user_id}")
            return jsonify({'error': 'Document not found'}), 404
            
        doc_info = uploaded_documents[user_id][doc_id]
        
        return jsonify({
            'doc_id': doc_id,
            'filename': doc_info['filename'],
            'content': doc_info['content'],
            'summary': doc_info['summary'],
            'user_type': doc_info.get('user_type', 'normal-user'),
            'upload_time': doc_info['upload_time']
        })
    except Exception as e:
        logger.error(f"Error retrieving document {doc_id}: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/process_query', methods=['POST'])
def process_query_endpoint():
    try:
        logger.info("Received request to /api/process_query endpoint")
        # Parse the incoming JSON data
        data = request.get_json()
        if not data:
            logger.warning("No JSON data in request")
            return jsonify({"error": "No data provided"}), 400
            
        query = data.get('query', '').lower()  # Normalize the query to lowercase
        user_type = data.get('user_type', 'normal-user')
        logger.info(f"Processing query: {query} for user type: {user_type}")
        
        # Check if the query type is related to a company
        if "company" in query:
            company_name = query.split("company")[-1].strip()  # Extract company name from the query
            response_data = handle_company_query(company_name)
            
            # Add detailed analysis for professional users
            if user_type == 'professional-user':
                response_data['detailed_analysis'] = True
                response_data['financial_metrics'] = {
                    'revenue': 'Placeholder for revenue data',
                    'growth': 'Placeholder for growth data',
                    'forecasts': 'Placeholder for forecast data'
                }
        
        # Check if the query is related to a sector
        elif "sector" in query:
            sector_name = query.split("sector")[-1].strip()  # Extract sector name from the query
            response_data = handle_sector_query(sector_name)
            
            # Add detailed analysis for professional users
            if user_type == 'professional-user':
                response_data['detailed_analysis'] = True
                response_data['sector_metrics'] = {
                    'market_size': 'Placeholder for market size data',
                    'trends': 'Placeholder for trend data',
                    'forecasts': 'Placeholder for forecast data'
                }

        # Default response for unrecognized queries
        else:
            logger.warning(f"Unrecognized query type: {query}")
            response_data = {
                "type": "unknown",
                "message": "Could not identify the query type."
            }

        logger.info(f"Returning response data: {response_data}")
        return jsonify({"response": json.dumps(response_data)})

    except Exception as e:
        logger.error(f"Error in process_query endpoint: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/reports/<report_id>', methods=['GET'])
def get_report(report_id):
    """Get a specific report"""
    try:
        user_type = request.args.get('user_type', 'normal-user')
        logger.info(f"Getting report with ID {report_id} for user type {user_type}")
        
        # Basic report for all users
        report = {
            'id': report_id,
            'title': 'Sample Report',
            'sections': [
                {
                    'title': 'Overview',
                    'content': 'This is an overview section of the report.'
                },
                {
                    'title': 'Key Metrics',
                    'content': 'This section contains key metrics and analysis.'
                },
                {
                    'title': 'Recommendations',
                    'content': 'Based on the analysis, here are the recommendations.'
                }
            ],
            'created_at': datetime.now().isoformat()
        }
        
        # Add additional sections for professional users
        if user_type == 'professional-user':
            report['sections'].extend([
                {
                    'title': 'Financial Analysis',
                    'content': 'Detailed financial analysis for professional users.'
                },
                {
                    'title': 'Market Forecasts',
                    'content': 'In-depth market forecasting data and projections.'
                },
                {
                    'title': 'Risk Assessment',
                    'content': 'Comprehensive risk assessment metrics and mitigation strategies.'
                }
            ])
            report['is_professional'] = True
        
        logger.info(f"Returning report data for ID {report_id}")
        return jsonify(report)
    except Exception as e:
        logger.error(f"Error in get_report endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    logger.info("Health check requested")
    return jsonify({
        'status': 'ok', 
        'version': '1.0.0',
        'timestamp': datetime.now().isoformat()
    })

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download a processed file"""
    try:
        logger.info(f"Download requested for file: {filename}")
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(file_path):
            logger.warning(f"File not found: {file_path}")
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        logger.error(f"Error downloading file: {str(e)}")
        return jsonify({'error': str(e)}), 500

# This route is important - it handles all other routes and forwards them to Next.js
@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def catch_all(path):
    """Serve the Next.js frontend for all non-API routes"""
    logger.info(f"Forwarding request to Next.js for path: /{path}")
    return jsonify({
        'message': 'This route should be handled by Next.js',
        'path': path
    })

@app.errorhandler(404)
def not_found(e):
    """Handle 404 errors"""
    if request.path.startswith('/api/'):
        logger.warning(f"API 404 error: {request.path}")
        return jsonify({'error': 'The requested API endpoint was not found'}), 404
    
    # For non-API routes, return a JSON response for the frontend to handle
    logger.info(f"Frontend route not found in backend: {request.path}")
    return jsonify({
        'error': 'Not found',
        'message': 'This route should be handled by Next.js',
        'path': request.path
    })

@app.errorhandler(500)
def internal_server_error(e):
    """Handle 500 errors"""
    logger.error(f"500 error: {str(e)}")
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    logger.info("Starting Flask application on port 5000")
    app.run(debug=True, host='0.0.0.0', port=5000)